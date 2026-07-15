from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class DocumentReadError(RuntimeError):
    pass


def read_document(path: str) -> str:
    file_path = Path(path)
    if not file_path.exists():
        raise DocumentReadError("文件不存在。")

    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise DocumentReadError("仅支持 .pdf、.docx、.txt 文件。")

    if suffix == ".pdf":
        return _read_pdf(file_path)
    if suffix == ".docx":
        return _read_docx(file_path)
    return _read_txt(file_path)


def _read_pdf(path: Path) -> str:
    chunks = []
    try:
        reader = PdfReader(str(path))
        for page in reader.pages:
            chunks.append(page.extract_text() or "")
    except Exception as exc:
        raise DocumentReadError(f"PDF 读取失败：{exc}") from exc

    return _normalize_text("\n".join(chunks))


def _read_docx(path: Path) -> str:
    try:
        document = Document(path)
    except Exception as exc:
        raise DocumentReadError(f"DOCX 读取失败：{exc}") from exc

    parts = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return _normalize_text("\n".join(parts))


def _read_txt(path: Path) -> str:
    encodings = ("utf-8", "utf-8-sig", "gbk", "gb18030")
    last_error = None
    for encoding in encodings:
        try:
            return _normalize_text(path.read_text(encoding=encoding))
        except UnicodeDecodeError as exc:
            last_error = exc
        except Exception as exc:
            raise DocumentReadError(f"TXT 读取失败：{exc}") from exc

    raise DocumentReadError(f"TXT 编码无法识别：{last_error}")


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    clean_lines = [line for line in lines if line]
    return "\n".join(clean_lines)
