from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


WBS_FILE_NAME = "工作分解结构_WBS.docx"
TASK_FILE_NAME = "任务清单.xlsx"


def export_deliverables(result: dict, output_dir: str | Path) -> dict[str, Path]:
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    wbs_path = output_path / WBS_FILE_NAME
    tasks_path = output_path / TASK_FILE_NAME

    _export_wbs_docx(result, wbs_path)
    _export_tasks_xlsx(result, tasks_path)
    return {"wbs": wbs_path, "tasks": tasks_path}


def save_deliverable(source: Path, target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, target)


def _export_wbs_docx(result: dict, path: Path) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)

    summary = result.get("project_summary", {})
    title = document.add_heading(summary.get("name") or "工作分解结构 WBS", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("一、项目概述", level=1)
    document.add_paragraph(str(summary.get("background", "待补充")))

    document.add_heading("二、项目目标", level=1)
    _add_bullets(document, summary.get("goals", []))

    document.add_heading("三、项目范围", level=1)
    _add_bullets(document, summary.get("scope", []))

    document.add_heading("四、工作分解结构", level=1)
    for item in result.get("wbs", []):
        document.add_paragraph(f"{item.get('id', '')} {item.get('name', '')}", style="List Number")
        for child in item.get("children", []) or []:
            text = f"{child.get('id', '')} {child.get('name', '')}"
            deliverable = child.get("deliverable")
            if deliverable:
                text += f" - 交付物：{deliverable}"
            paragraph = document.add_paragraph(text, style="List Bullet")
            paragraph.paragraph_format.left_indent = Pt(18)

    document.add_heading("五、假设与风险", level=1)
    _add_bullets(document, summary.get("assumptions", []))

    document.add_heading("六、任务概览", level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["任务编号", "WBS编号", "阶段", "任务名称", "负责人角色"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for task in result.get("tasks", []):
        cells = table.add_row().cells
        cells[0].text = str(task.get("task_id", ""))
        cells[1].text = str(task.get("wbs_id", ""))
        cells[2].text = str(task.get("phase", ""))
        cells[3].text = str(task.get("task_name", ""))
        cells[4].text = str(task.get("owner_role", ""))

    document.save(path)


def _export_tasks_xlsx(result: dict, path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "任务清单"

    headers = [
        ("task_id", "任务编号"),
        ("wbs_id", "WBS编号"),
        ("phase", "阶段"),
        ("task_name", "任务名称"),
        ("description", "任务说明"),
        ("owner_role", "负责人角色"),
        ("input", "输入物"),
        ("output", "输出物"),
        ("priority", "优先级"),
        ("estimated_duration", "预计工期"),
        ("dependency", "前置依赖"),
        ("acceptance_criteria", "验收标准"),
    ]

    header_fill = PatternFill("solid", fgColor="D9EAF7")
    header_font = Font(bold=True)
    for column, (_, label) in enumerate(headers, start=1):
        cell = sheet.cell(row=1, column=column, value=label)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for row_index, task in enumerate(result.get("tasks", []), start=2):
        for column, (key, _) in enumerate(headers, start=1):
            cell = sheet.cell(row=row_index, column=column, value=str(task.get(key, "")))
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    widths = [12, 12, 14, 24, 42, 16, 24, 24, 10, 12, 16, 42]
    for column, width in enumerate(widths, start=1):
        sheet.column_dimensions[get_column_letter(column)].width = width

    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    workbook.save(path)


def _add_bullets(document: Document, items: list) -> None:
    if not items:
        document.add_paragraph("待补充")
        return
    for item in items:
        document.add_paragraph(str(item), style="List Bullet")
